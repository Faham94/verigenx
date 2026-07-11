"""
Document Ingestion for SpecMind
Supports TXT, PDF (PyMuPDF), DOCX (tables + paragraphs),
IP-XACT XML, SystemRDL, and incremental hash-based caching.

Fixes applied:
  - Bug #6:  PDF close() called before len(doc) → fixed order
  - Bug #7:  DOCX ignored tables → now extracts all table cells
  - Bug #10: IP-XACT / SystemRDL XML formats now supported
  - Bug #8:  Incremental update cache via SHA-256 file hashing
"""
import os
import re
import json
import hashlib
from pathlib import Path
from VeriGenX.config import CACHE_DIR


class DocumentIngestor:

    # ------------------------------------------------------------------ #
    #  Public entry point                                                  #
    # ------------------------------------------------------------------ #
    def ingest(self, file_path: str) -> dict:
        """
        Ingest a specification file and return structured text + metadata.
        Uses SHA-256 hash caching for incremental updates — unchanged files
        return the cached result without re-parsing.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # --- Incremental caching (Bug #8 fix) ---
        file_hash = self._hash_file(file_path)
        cached    = self._load_cache(file_path, file_hash)
        if cached is not None:
            print(f"  [Cache HIT] Using cached result for {os.path.basename(file_path)}")
            return cached

        ext    = Path(file_path).suffix.lower()
        result = self._dispatch(file_path, ext)

        self._save_cache(file_path, file_hash, result)
        return result

    # ------------------------------------------------------------------ #
    #  Format dispatching                                                  #
    # ------------------------------------------------------------------ #
    def _dispatch(self, file_path: str, ext: str) -> dict:
        if ext == ".txt":
            return self._ingest_txt(file_path)
        elif ext == ".pdf":
            return self._ingest_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return self._ingest_docx(file_path)
        elif ext == ".xml":
            return self._ingest_ipxact(file_path)
        elif ext == ".rdl":
            return self._ingest_systemrdl(file_path)
        else:
            raise ValueError(
                f"Unsupported file format: {ext}. "
                "Supported: .txt .pdf .docx .xml (IP-XACT) .rdl (SystemRDL)"
            )

    # ------------------------------------------------------------------ #
    #  Parsers                                                             #
    # ------------------------------------------------------------------ #
    def _ingest_txt(self, file_path: str) -> dict:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        return {
            "text":      text,
            "file_type": "txt",
            "metadata":  {"title": os.path.basename(file_path), "chars": len(text)},
        }

    def _ingest_pdf(self, file_path: str) -> dict:
        """
        Bug #6 fix: capture len(doc) BEFORE close() is called.
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF not installed. Run: pip install PyMuPDF")

        doc        = fitz.open(file_path)
        page_count = len(doc)           # FIXED: read length before closing
        text       = ""
        for page in doc:
            text += page.get_text()
        doc.close()

        return {
            "text":      text,
            "file_type": "pdf",
            "metadata":  {
                "title": os.path.basename(file_path),
                "pages": page_count,
                "chars": len(text),
            },
        }

    def _ingest_docx(self, file_path: str) -> dict:
        """
        Bug #7 fix: extract both paragraphs AND all table cells.
        Register maps and signal tables in DOCX are now captured.
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc   = Document(file_path)
        parts = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Tables — Bug #7 fix
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append("\t".join(cells))

        text = "\n".join(parts)
        return {
            "text":      text,
            "file_type": "docx",
            "metadata":  {"title": os.path.basename(file_path), "chars": len(text)},
        }

    def _ingest_ipxact(self, file_path: str) -> dict:
        """
        Bug #10 fix: parse IP-XACT XML format.
        Extracts component name, bus interfaces, registers, and fields.
        """
        import xml.etree.ElementTree as ET

        tree = ET.parse(file_path)
        root = tree.getroot()

        # IP-XACT uses namespaces — strip them for simple parsing
        def strip_ns(tag: str) -> str:
            return re.sub(r"\{[^}]*\}", "", tag)

        def iter_text(element, depth=0) -> list:
            lines = []
            tag   = strip_ns(element.tag)
            text  = (element.text or "").strip()
            if text:
                lines.append(f"{'  ' * depth}{tag}: {text}")
            for child in element:
                lines.extend(iter_text(child, depth + 1))
            return lines

        lines = iter_text(root)
        text  = "\n".join(lines)

        return {
            "text":      text,
            "file_type": "ipxact_xml",
            "metadata":  {
                "title":  os.path.basename(file_path),
                "format": "IP-XACT",
                "chars":  len(text),
            },
        }

    def _ingest_systemrdl(self, file_path: str) -> dict:
        """
        Bug #10 fix: parse SystemRDL (.rdl) as structured text.
        Extracts component declarations, registers, fields, and properties.
        """
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read()

        # Extract structural elements via regex
        parts = []

        # Component / addrmap / regfile declarations
        for m in re.finditer(r"(addrmap|regfile|component|reg|field)\s+(\w+)", raw):
            parts.append(f"{m.group(1).upper()}: {m.group(2)}")

        # Properties: name = value
        for m in re.finditer(r"(\w+)\s*=\s*([^;]+);", raw):
            parts.append(f"  {m.group(1)} = {m.group(2).strip()}")

        text = "\n".join(parts) if parts else raw
        return {
            "text":      text,
            "file_type": "systemrdl",
            "metadata":  {
                "title":  os.path.basename(file_path),
                "format": "SystemRDL",
                "chars":  len(text),
            },
        }

    # ------------------------------------------------------------------ #
    #  Incremental caching (Bug #8)                                        #
    # ------------------------------------------------------------------ #
    @staticmethod
    def _hash_file(file_path: str) -> str:
        h = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(65536), b""):
                h.update(chunk)
        return h.hexdigest()

    @staticmethod
    def _cache_path(file_path: str) -> str:
        os.makedirs(CACHE_DIR, exist_ok=True)
        safe_name = Path(file_path).stem + ".cache.json"
        return os.path.join(CACHE_DIR, safe_name)

    def _load_cache(self, file_path: str, file_hash: str):
        cp = self._cache_path(file_path)
        if not os.path.exists(cp):
            return None
        try:
            with open(cp, "r", encoding="utf-8") as f:
                entry = json.load(f)
            if entry.get("hash") == file_hash:
                return entry.get("result")
        except Exception:
            pass
        return None

    def _save_cache(self, file_path: str, file_hash: str, result: dict):
        cp = self._cache_path(file_path)
        try:
            with open(cp, "w", encoding="utf-8") as f:
                json.dump({"hash": file_hash, "result": result}, f, indent=2)
        except Exception:
            pass
